import io
from pypdf import PdfReader
from docx import Document

class FileParser:
    @staticmethod
    def __extract_text_from_pdf(pdf_bytes):
        try:
            pdf_file_object = io.BytesIO(pdf_bytes)
            reader = PdfReader(pdf_file_object)

            extracted_text = ""
            for page in reader.pages:
                extracted_text += page.extract_text() 

            return extracted_text
        except Exception as e:
            print(f"Error extracting text from PDF bytes: {e}")
            return ""

    @staticmethod
    def __extract_text_from_docx(docx_bytes):
        try:
            docx_file_object = io.BytesIO(docx_bytes)
            doc = Document(docx_file_object)

            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())

            return "\n".join(paragraphs)
        except Exception as e:
            print(f"Error extracting text from DOCX bytes: {e}")
            return ""

    @staticmethod   
    def extract_text(file_bytes: bytes, file_extention: str) -> str:
        file_extention = file_extention.lower()

        if file_extention == "pdf":
            return FileParser.__extract_text_from_pdf(file_bytes)
        elif file_extention in ("docx", "doc"):
            return FileParser.__extract_text_from_docx(file_bytes)
        else:
            return ""